# -*- coding: utf-8 -*-
"""
中興大學生命科學系研究所博士班專題討論 — 摘要檢查器
============================================================
依「摘要繕寫模式」（114-1 學年度老師更新版）+ Zoological Studies 引用格式

用法：
    python check_abstract.py 你的摘要.docx
或在 Windows 上將 docx 拖放到「拖放這裡檢查.bat」

規則來源：
    1. 啟動時嘗試從 rules_remote_url.txt 指定的 URL 拉最新規則
    2. 拉不到就用本地 rules.yaml
    3. 規則檔過期會強制警告

GitHub: https://github.com/gjj22622/nchu-lifeSciences-seminar-abstract-checker

授權：MIT License (見 LICENSE)
作者：鐘基啓（博二，學號 8113052008）2026 年留給後輩
"""
from __future__ import annotations
import re
import sys
import io
import json
import urllib.request
import urllib.error
from datetime import date, datetime
from pathlib import Path

# ----- Windows 終端中文不亂碼 -----
if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    except AttributeError:
        pass

# ----- 依賴檢查 -----
try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("錯誤：尚未安裝 python-docx 套件。")
    print("請執行：  pip install -r requirements.txt")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("錯誤：尚未安裝 pyyaml 套件。")
    print("請執行：  pip install -r requirements.txt")
    sys.exit(1)

# ----- 報告等級 -----
PASS = "[通過]"
WARN = "[提醒]"
FAIL = "[必改]"
INFO = "[資訊]"

# ----- 路徑常數 -----
SCRIPT_DIR = Path(__file__).parent
LOCAL_RULES_PATH = SCRIPT_DIR / "rules.yaml"
REMOTE_URL_PATH = SCRIPT_DIR / "rules_remote_url.txt"
CACHE_PATH = SCRIPT_DIR / ".rules_cache.yaml"
FETCH_TIMEOUT = 3  # seconds


def _looks_like_latin_binomial_token(text: str) -> bool:
    """偵測 run 內容是否為拉丁學名片段（避免誤判學名為違規斜體）。
    支援：屬名（首大寫其餘小寫）、種小名（全小寫）、縮寫屬名（X.）、
    縮寫屬名+種名（X. ssp）、屬名+種名整段、空白片段。
    """
    text = text.strip()
    if not text:
        return True
    # 屬名單字：Pinus, Metasequoia
    if re.fullmatch(r"[A-Z][a-z]{2,}", text):
        return True
    # 種小名：pinaster, glyptostroboides
    if re.fullmatch(r"[a-z][a-z\-]{2,}", text):
        return True
    # 縮寫屬名：P., M.
    if re.fullmatch(r"[A-Z]\.", text):
        return True
    # 縮寫屬名 + 空白 + 種小名：P. radiata
    if re.fullmatch(r"[A-Z]\.\s+[a-z][a-z\-]+", text):
        return True
    # 完整屬名 + 空白 + 種小名：Pinus pinaster
    if re.fullmatch(r"[A-Z][a-z]+\s+[a-z][a-z\-]+", text):
        return True
    # 包含三名法（subsp./var.）的學名
    if re.fullmatch(r"[A-Z][a-z]+\s+[a-z\-]+(\s+(subsp|var|f)\.\s+[a-z\-]+)?", text):
        return True
    return False


# =============================================================================
# 規則載入器（遠端優先 + 本地 fallback + 過期警告）
# =============================================================================
class RuleLoader:
    def __init__(self):
        self.rules = None
        self.source = None  # 'remote' / 'cache' / 'local'
        self.warnings: list[str] = []

    def load(self) -> dict:
        if self._try_remote():
            self.source = "remote"
            return self.rules
        if self._try_cache():
            self.source = "cache"
            self.warnings.append("使用本地快取規則（無法連線取得最新版）")
            return self.rules
        self._load_local()
        self.source = "local"
        self.warnings.append("使用內建預設規則（無快取、無網路）")
        return self.rules

    def _try_remote(self) -> bool:
        if not REMOTE_URL_PATH.exists():
            return False
        try:
            url = REMOTE_URL_PATH.read_text(encoding="utf-8").strip()
            if not url or url.startswith("#"):
                return False
            req = urllib.request.Request(
                url,
                headers={"User-Agent": "nchu-lifeSciences-abstract-checker/1.0"},
            )
            with urllib.request.urlopen(req, timeout=FETCH_TIMEOUT) as resp:
                content = resp.read().decode("utf-8")
            self.rules = yaml.safe_load(content)
            CACHE_PATH.write_text(content, encoding="utf-8")
            return True
        except (urllib.error.URLError, TimeoutError, OSError, yaml.YAMLError) as e:
            self.warnings.append(f"遠端規則拉取失敗：{type(e).__name__}")
            return False

    def _try_cache(self) -> bool:
        if not CACHE_PATH.exists():
            return False
        try:
            self.rules = yaml.safe_load(CACHE_PATH.read_text(encoding="utf-8"))
            return True
        except (OSError, yaml.YAMLError):
            return False

    def _load_local(self):
        if not LOCAL_RULES_PATH.exists():
            print(f"嚴重錯誤：找不到本地規則檔 {LOCAL_RULES_PATH}")
            sys.exit(2)
        self.rules = yaml.safe_load(LOCAL_RULES_PATH.read_text(encoding="utf-8"))

    def expiry_warning(self) -> str | None:
        """若規則檔已過期，回傳警告訊息；否則 None。"""
        try:
            valid_until = datetime.strptime(self.rules["valid_until"], "%Y-%m-%d").date()
        except (KeyError, ValueError):
            return None
        today = date.today()
        if today > valid_until:
            days_over = (today - valid_until).days
            maintainer = self.rules.get("maintainer", {})
            return (
                f"規則檔已超過有效日期 {valid_until}（已逾期 {days_over} 天）。"
                f"請向學長姐／系辦／黃老師確認最新規範後再交件。"
                f"原作者：{maintainer.get('name', '?')} <{maintainer.get('email', '?')}>"
            )
        return None


# =============================================================================
# 摘要檢查器
# =============================================================================
class AbstractChecker:
    def __init__(self, path: Path, rules: dict):
        self.path = path
        self.rules = rules
        self.doc = Document(path)
        self.issues: list[tuple[str, str, str]] = []
        self.title_paragraphs: list = []
        self.speaker_paragraphs: list = []
        self.abstract_paragraph = None
        self.reference_paragraphs: list = []

    def add(self, level: str, category: str, message: str):
        self.issues.append((level, category, message))

    # --- 1. 識別段落 ---
    def classify_paragraphs(self):
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if re.match(r"^\d+\.\s", text):
                self.reference_paragraphs.append(p)
                continue
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and any(r.bold for r in p.runs if r.text):
                self.title_paragraphs.append(p)
                continue
            if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                self.speaker_paragraphs.append(p)
                continue
            if p.paragraph_format.first_line_indent and len(text) > 100:
                self.abstract_paragraph = p
                continue
            if not self.abstract_paragraph and len(text) > 200:
                self.abstract_paragraph = p

    # --- 2. 邊界檢查 ---
    def check_margins(self):
        target = self.rules["format"]["margin_cm"]
        tolerance = self.rules["format"]["margin_tolerance_cm"]
        sec = self.doc.sections[0]
        margins = {"上": sec.top_margin, "下": sec.bottom_margin,
                   "左": sec.left_margin, "右": sec.right_margin}
        all_ok = True
        for name, val in margins.items():
            cm = val / Cm(1)
            if abs(cm - target) > tolerance:
                self.add(FAIL, "版面", f"{name}邊界 = {cm:.2f} cm，應為 {target} cm")
                all_ok = False
        if all_ok:
            self.add(PASS, "版面", f"邊界 {target} cm × 4 邊")

    # --- 3. 標題檢查 ---
    def check_titles(self):
        cfg = self.rules["format"]["title"]
        if len(self.title_paragraphs) < 2:
            self.add(FAIL, "標題",
                     f"應有中、英文標題各一行；偵測到 {len(self.title_paragraphs)} 行置中粗體段落")
            return
        target_pt = cfg["font_size_pt"]
        for tag, p in [("中文標題", self.title_paragraphs[0]),
                        ("英文標題", self.title_paragraphs[1])]:
            sizes = [r.font.size for r in p.runs if r.text and r.font.size]
            if sizes:
                pt = sizes[0] / Pt(1)
                if abs(pt - target_pt) > 0.5:
                    self.add(FAIL, "標題", f"{tag}字級 = {pt:.0f} pt，應為 {target_pt} pt")
                else:
                    self.add(PASS, "標題", f"{tag}字級 {target_pt} pt 粗體置中")
            else:
                self.add(WARN, "標題", f"{tag}字級無法偵測（可能用了預設字級）")
        # 冒號偵測
        forbidden = self.rules["style"]["forbidden_in_title"]
        for tag, p in [("中文", self.title_paragraphs[0]),
                        ("英文", self.title_paragraphs[1])]:
            for f in forbidden:
                if f in p.text:
                    self.add(WARN, "標題", f"{tag}標題含 `{f}`，易被視為 AI 味；建議改命題式單句")

    # --- 4. 講者資訊檢查 ---
    def check_speaker_info(self):
        cfg = self.rules["format"]["speaker_info"]
        text_all = "\n".join(p.text for p in self.speaker_paragraphs)
        if not text_all.strip():
            self.add(FAIL, "講者資訊", "未偵測到右下角講者資訊（姓名、學號、日期）")
            return
        ok_name = bool(re.search(r"[一-鿿]{2,4}", text_all))
        ok_id = bool(re.search(cfg["student_id_pattern"], text_all))
        date_pat = "|".join(cfg["date_patterns"])
        ok_date = bool(re.search(date_pat, text_all))
        if cfg.get("require_name") and not ok_name:
            self.add(WARN, "講者資訊", "未偵測到中文姓名")
        if cfg.get("require_student_id") and not ok_id:
            self.add(FAIL, "講者資訊", "未偵測到學號（8-12 位數字）")
        if cfg.get("require_date") and not ok_date:
            self.add(WARN, "講者資訊", "未偵測到報告日期；建議格式 MM/DD/YYYY")
        if ok_name and ok_id and ok_date:
            self.add(PASS, "講者資訊", "姓名、學號、日期皆齊全且靠右對齊")

    # --- 5. 摘要正文檢查 ---
    def check_abstract_body(self):
        cfg = self.rules["format"]["abstract_body"]
        if not self.abstract_paragraph:
            self.add(FAIL, "摘要正文", "找不到摘要正文段落（應為單一長段、首行縮入、左右對齊）")
            return
        p = self.abstract_paragraph
        text = p.text.strip()
        # 字數
        cn_chars = len(re.findall(r"[一-鿿]", text))
        en_words = len(re.findall(r"\b[A-Za-z][A-Za-z\-]*\b", text))
        if cn_chars > en_words * 3:
            count, unit = cn_chars, "字"
            mn, mx = cfg["word_count_min_cn"], cfg["word_count_max_cn"]
        else:
            count, unit = en_words, "詞"
            mn, mx = cfg["word_count_min_en"], cfg["word_count_max_en"]
        tol_pct = cfg.get("word_count_tolerance_pct", 0)
        mx_tol = int(mx * (1 + tol_pct / 100))
        if mn <= count <= mx:
            self.add(PASS, "摘要字數", f"{count} {unit}（規範 {mn}–{mx}）")
        elif mx < count <= mx_tol:
            self.add(WARN, "摘要字數", f"{count} {unit}（超過 {mx} 但在 {tol_pct}% 容忍內 ≤{mx_tol}；老師通常接受）")
        elif count < mn:
            self.add(FAIL, "摘要字數", f"{count} {unit}，少於規範下限 {mn}")
        else:
            self.add(FAIL, "摘要字數", f"{count} {unit}，超過 {tol_pct}% 容忍上限 {mx_tol}")
        # 對齊
        if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY and cfg.get("alignment") == "justify":
            self.add(WARN, "摘要正文", "未設左右對齊（justify）")
        # 首行縮入
        target_indent = cfg["first_line_indent_cm"]
        fli = p.paragraph_format.first_line_indent
        if not fli or abs(fli / Cm(1) - target_indent) > 0.1:
            actual = f"{fli/Cm(1):.2f} cm" if fli else "未設"
            self.add(FAIL, "摘要正文", f"首行縮入 = {actual}，應為 {target_indent} cm")
        else:
            self.add(PASS, "摘要正文", f"首行縮入 {target_indent} cm")
        # 字級
        target_pt = cfg["font_size_pt"]
        sizes = [r.font.size for r in p.runs if r.text and r.font.size]
        if sizes:
            pt = sizes[0] / Pt(1)
            if abs(pt - target_pt) > 0.5:
                self.add(FAIL, "摘要正文", f"字級 = {pt:.0f} pt，應為 {target_pt} pt")
            else:
                self.add(PASS, "摘要正文", f"字級 {target_pt} pt")
        # 粗體
        bold_runs = [r for r in p.runs if r.bold and r.text.strip()]
        if bold_runs and not all(r.bold for r in p.runs if r.text.strip()):
            self.add(FAIL, "摘要正文", "正文出現部分粗體（規範：不用粗體）")
        # 斜體（學名為合法斜體，需濾除避免 false positive）
        italic_runs = [r for r in p.runs if r.italic and r.text.strip()]
        non_latin = [r for r in italic_runs if not _looks_like_latin_binomial_token(r.text)]
        if non_latin:
            italic_text = " | ".join(r.text.strip()[:30] for r in non_latin)
            self.add(WARN, "摘要正文", f"含斜體段落（規範：除學名外不用斜體）：{italic_text}")
        # AI 味禁用詞
        bad = self.rules["style"]["ai_bad_words_cn"] + self.rules["style"]["ai_bad_words_en"]
        for w in bad:
            if w.lower() in text.lower():
                self.add(WARN, "AI 味", f"含潛在 AI 味詞：「{w}」")
        # 新名詞定義
        nt = self.rules["style"].get("new_term_definition", {})
        if nt.get("enabled"):
            for term in nt.get("terms", []):
                if term.lower() in text.lower():
                    keys = nt.get("must_contain_any_of", [])
                    if not any(k.lower() in text.lower() for k in keys):
                        self.add(WARN, "新名詞定義",
                                 f"摘要出現「{term}」但未見定義字樣（應包含 {keys} 之一）")

    # --- 6. 引用文獻檢查 ---
    def check_references(self):
        cfg = self.rules["reference"]
        n = len(self.reference_paragraphs)
        if n == 0:
            self.add(FAIL, "引用文獻", "未偵測到引用清單（以「1. 」「2. 」開頭的段落）")
            return
        self.add(INFO, "引用文獻", f"偵測到 {n} 條引用")

        # 主要參考粗體
        bold_refs = [p for p in self.reference_paragraphs
                     if p.runs and all(r.bold for r in p.runs if r.text.strip())]
        if cfg.get("must_have_main_reference_bold") and not bold_refs:
            self.add(WARN, "主要參考", "未偵測到任何整條粗體的「主要參考報告」")
        elif bold_refs:
            self.add(PASS, "主要參考", f"{len(bold_refs)} 條主要參考已粗體標示")

        # Last name 排序
        if cfg.get("must_be_alphabetical_by_lastname"):
            first_letters = []
            for p in self.reference_paragraphs:
                m = re.match(r"^\d+\.\s+([A-Za-zÀ-ÿÀ-ſ]+)", p.text.strip())
                if m:
                    first_letters.append(m.group(1).upper())
            if first_letters == sorted(first_letters):
                self.add(PASS, "引用排序", "依 Last name 字母順序")
            else:
                self.add(FAIL, "引用排序", f"未依字母順序排：{first_letters}")

        # 個別逐條檢查
        et_al_threshold = cfg.get("et_al_threshold", 4)
        et_al_list_count = cfg.get("et_al_list_count", 3)
        for idx, p in enumerate(self.reference_paragraphs, 1):
            text = p.text.strip()
            short = text[:60] + ("..." if len(text) > 60 else "")
            # DOI
            if not cfg.get("doi_should_appear") and ("doi:" in text.lower() or "doi.org" in text.lower()):
                self.add(WARN, f"引用#{idx}", f"含 DOI（114-1 規範範例皆無 DOI）：{short}")
            # 期刊縮寫帶句點
            if cfg.get("journal_abbrev_no_dot"):
                if re.search(r"\.\s+([A-Z][a-z]+\.\s+)+[A-Z][a-z]*\.?\s+\d+:", text):
                    self.add(FAIL, f"引用#{idx}",
                             f"期刊縮寫帶句點（規範：縮寫無句點）：{short}")
            # et al. 規則（Zool Stud 真實 article 三種做法並存，預設不檢查）
            et_al_sev = cfg.get("et_al_severity", "warn")
            if et_al_sev != "off":
                m = re.match(r"^\d+\.\s+(.+?)\.\s+\d{4}\.", text)
                if m:
                    authors_str = m.group(1)
                    has_et_al = "et al" in authors_str.lower()
                    level = FAIL if et_al_sev == "fail" else WARN
                    if has_et_al:
                        listed = len([a for a in authors_str.split(",") if "et al" not in a.lower()])
                        if listed != et_al_list_count:
                            self.add(level, f"引用#{idx}",
                                     f"使用 et al. 但列出 {listed} 位（規範：列前 {et_al_list_count} 位 + et al.）")
                    else:
                        listed = len(authors_str.split(","))
                        if listed >= et_al_threshold:
                            self.add(level, f"引用#{idx}",
                                     f"作者數 = {listed}（≥{et_al_threshold} 應用「列前 {et_al_list_count} 位 + et al.」）")
            # 卷號粗體
            if cfg.get("volume_must_be_bold") and p not in bold_refs:
                vol_match = re.search(r"\s(\d+):\d", text)
                if vol_match:
                    vol_str = vol_match.group(1) + ":"
                    found = any(vol_str in r.text and r.bold for r in p.runs)
                    if not found:
                        self.add(FAIL, f"引用#{idx}", f"卷號 `{vol_str}` 未粗體：{short}")
            # 頁碼 hyphen
            sev = cfg.get("hyphen_in_pages_severity", "warn")
            if sev != "off":
                if re.search(r"\d+\s*[-]\s*\d+", text) and not re.search(r"\d+\s*[–]\s*\d+", text):
                    level = FAIL if sev == "fail" else WARN
                    self.add(level, f"引用#{idx}",
                             f"頁碼疑似用 hyphen（-）而非 en-dash（–）：{short}")
            # title case 偵測
            sev = cfg.get("title_case_severity", "warn")
            if sev != "off":
                title_match = re.search(r"\.\s+\d{4}\.\s+(.+?)\.\s+[A-Z][a-z]+\s+\d+:", text)
                if title_match:
                    title = title_match.group(1)
                    words = title.split()
                    if len(words) > 5:
                        exceptions = set(s.lower() for s in cfg.get("sentence_case_exceptions", []))
                        cap = sum(1 for w in words[1:]
                                  if w[:1].isupper() and not w.isupper() and len(w) > 3
                                  and w.lower() not in exceptions)
                        if cap / len(words) > 0.3:
                            level = FAIL if sev == "fail" else WARN
                            self.add(level, f"引用#{idx}",
                                     f"標題疑似 title case（規範：sentence case）：{title[:50]}...")

    # --- 7. 一頁估算 ---
    def check_one_page(self):
        threshold = self.rules["one_page"]["paragraph_count_warn_threshold"]
        n = sum(1 for p in self.doc.paragraphs if p.text.strip())
        if n > threshold:
            self.add(WARN, "頁數", f"段落數 = {n}（>{threshold}），A4 一頁可能不夠；超頁需雙面印")

    def run_all(self):
        self.classify_paragraphs()
        self.check_margins()
        self.check_titles()
        self.check_speaker_info()
        self.check_abstract_body()
        self.check_references()
        self.check_one_page()

    def report(self, rule_loader: RuleLoader) -> int:
        levels = {PASS: 0, WARN: 0, FAIL: 0, INFO: 0}
        for level, _, _ in self.issues:
            levels[level] += 1
        print("=" * 72)
        print("中興生科所專討摘要檢查報告")
        print(f"檔案：{self.path.name}")
        print(f"規則：{self.rules.get('spec_version', '?')} ／ "
              f"v{self.rules.get('version', '?')} ／ "
              f"來源 {rule_loader.source}")
        # 過期警告
        expiry = rule_loader.expiry_warning()
        if expiry:
            print(f"\n  ⚠️ {expiry}\n")
        for w in rule_loader.warnings:
            print(f"  · {w}")
        print("=" * 72)
        order = {FAIL: 0, WARN: 1, INFO: 2, PASS: 3}
        for level, cat, msg in sorted(self.issues, key=lambda x: order[x[0]]):
            print(f"  {level} [{cat}] {msg}")
        print("-" * 72)
        print(f"總計：{levels[FAIL]} 必改 ／ {levels[WARN]} 提醒 ／ "
              f"{levels[PASS]} 通過 ／ {levels[INFO]} 資訊")
        if levels[FAIL] == 0 and levels[WARN] == 0:
            print("\n  >>> 太棒了！未偵測到任何問題，可直接交件。")
        elif levels[FAIL] == 0:
            print(f"\n  >>> 可交件，但有 {levels[WARN]} 項建議改善。")
        else:
            print(f"\n  >>> 有 {levels[FAIL]} 項必改錯誤，請修正後再跑一次檢查。")
        print("=" * 72)
        return 0 if levels[FAIL] == 0 else 1


def main() -> int:
    if len(sys.argv) < 2:
        print(__doc__)
        return 1
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"錯誤：找不到檔案 {path}")
        return 1
    if path.suffix.lower() != ".docx":
        print(f"錯誤：僅支援 .docx 檔案（你給的是 {path.suffix}）")
        return 1
    loader = RuleLoader()
    rules = loader.load()
    try:
        checker = AbstractChecker(path, rules)
        checker.run_all()
        return checker.report(loader)
    except Exception as e:
        print(f"\n檢查器執行錯誤：{type(e).__name__}: {e}")
        print("請改開 Word 對照「規範速查.md」手動檢查。")
        return 2


if __name__ == "__main__":
    sys.exit(main())
